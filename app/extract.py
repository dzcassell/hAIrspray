"""File extraction for the Prompt panel's attach feature.

Three categories of input:

* **Inline-text formats** (.txt, .md, .csv, .json, .log, .xml, .yml,
  .yaml, .html, .htm, .tsv, .ini, .conf, .py, .js, .sql) — read as
  UTF-8, fall back to latin-1 on decode error. No parsing; the
  bytes-to-text round trip is the whole job.

* **Structured documents** (.pdf, .docx, .xlsx) — third-party
  libraries pull text out. pypdf for PDFs, python-docx for Word
  documents, openpyxl for spreadsheets. Each library is pure-Python
  with no native dependencies, so no Dockerfile changes are needed.

* **Anything else** — rejected with a clear error. We don't try to
  guess at unknown extensions; the operator picks a different file.

Extraction caps:

* The raw upload is capped at 10 MB by the web layer before we get
  here. This module assumes whatever it's handed is within that.
* Extracted text is capped at MAX_EXTRACT_CHARS (~20K). Beyond that,
  we truncate and append a notice — most provider context windows
  will choke at substantially more than this anyway, and a 20K-char
  attachment is already a more thorough DLP test than 99% of real
  user uploads.

Public API:

* extract(filename, content_bytes) -> ExtractResult

ExtractResult carries text, char_count, truncated flag, source-format
hint, and a one-line human-readable summary suitable for showing
in the chip UI.
"""
from __future__ import annotations

import io
import logging
from dataclasses import dataclass
from pathlib import Path

log = logging.getLogger(__name__)

# Hard cap on extracted text we'll send to a model. Above this we
# truncate and tell the model the doc was longer.
MAX_EXTRACT_CHARS = 20_000

# Inline-text extensions (no parser; just a decode pass).
INLINE_TEXT_EXTS = {
    ".txt", ".md", ".markdown", ".csv", ".tsv", ".json", ".log",
    ".xml", ".yml", ".yaml", ".html", ".htm", ".ini", ".conf",
    ".cfg", ".toml", ".env",
    # Source code — frequently uploaded for "explain this" prompts,
    # which is exactly the kind of DLP-relevant flow a SASE wants
    # to classify and inspect.
    ".py", ".js", ".ts", ".jsx", ".tsx", ".sh", ".bash", ".zsh",
    ".sql", ".rb", ".go", ".rs", ".java", ".kt", ".c", ".cc", ".cpp",
    ".h", ".hpp", ".cs", ".php", ".pl", ".lua", ".swift",
}

# Document formats we know how to extract from.
DOC_EXTS = {".pdf", ".docx", ".xlsx"}

# Everything we'll accept (text + docs).
SUPPORTED_EXTS = INLINE_TEXT_EXTS | DOC_EXTS


@dataclass(frozen=True)
class ExtractResult:
    """Outcome of extracting one upload."""
    text: str
    char_count: int       # length of the extracted text BEFORE truncation
    truncated: bool       # was the text shortened to MAX_EXTRACT_CHARS?
    source_kind: str      # "text", "pdf", "docx", "xlsx"
    summary: str          # one-line human description for the UI chip


# ---------------------------------------------------------------------------
# Inline text — UTF-8 with latin-1 fallback
# ---------------------------------------------------------------------------

def _decode_text_bytes(data: bytes) -> str:
    try:
        return data.decode("utf-8")
    except UnicodeDecodeError:
        # latin-1 is a "never fails" 1:1 byte-to-char mapping, so it
        # always decodes — the result is just gibberish if the file
        # was actually some other encoding. Better than refusing the
        # upload outright.
        return data.decode("latin-1", errors="replace")


def _extract_inline_text(filename: str, data: bytes) -> ExtractResult:
    raw = _decode_text_bytes(data)
    full_chars = len(raw)
    truncated = full_chars > MAX_EXTRACT_CHARS
    text = raw[:MAX_EXTRACT_CHARS] if truncated else raw
    if truncated:
        text += (
            f"\n\n[...truncated, full file was {full_chars:,} characters; "
            f"only the first {MAX_EXTRACT_CHARS:,} are included]"
        )
    return ExtractResult(
        text=text,
        char_count=full_chars,
        truncated=truncated,
        source_kind="text",
        summary=f"{full_chars:,} chars" + (" (truncated)" if truncated else ""),
    )


# ---------------------------------------------------------------------------
# PDF — pypdf
# ---------------------------------------------------------------------------

def _extract_pdf(filename: str, data: bytes) -> ExtractResult:
    # pypdf is heavy enough to not import at module load — defer.
    import pypdf  # type: ignore[import-not-found]

    reader = pypdf.PdfReader(io.BytesIO(data))
    page_count = len(reader.pages)
    parts: list[str] = []
    for i, page in enumerate(reader.pages):
        try:
            t = page.extract_text() or ""
        except Exception as e:  # noqa: BLE001 — pypdf raises a zoo
            log.warning("pdf_page_extract_failed page=%d err=%s", i, e)
            t = ""
        if t.strip():
            parts.append(f"--- Page {i + 1} ---\n{t.strip()}")

    if not parts:
        # Common when a PDF is scanned-image-only: pypdf finds no text
        # because there's no text layer. Tell the operator clearly.
        return ExtractResult(
            text="",
            char_count=0,
            truncated=False,
            source_kind="pdf",
            summary=f"PDF, {page_count} pages, no text layer (scanned image?)",
        )

    raw = "\n\n".join(parts)
    full_chars = len(raw)
    truncated = full_chars > MAX_EXTRACT_CHARS
    text = raw[:MAX_EXTRACT_CHARS] if truncated else raw
    if truncated:
        text += (
            f"\n\n[...truncated, full PDF was {full_chars:,} characters; "
            f"showing first {MAX_EXTRACT_CHARS:,}]"
        )
    return ExtractResult(
        text=text,
        char_count=full_chars,
        truncated=truncated,
        source_kind="pdf",
        summary=(f"PDF, {page_count} pages, {full_chars:,} chars"
                 + (" (truncated)" if truncated else "")),
    )


# ---------------------------------------------------------------------------
# DOCX — python-docx
# ---------------------------------------------------------------------------

def _extract_docx(filename: str, data: bytes) -> ExtractResult:
    import docx  # type: ignore[import-not-found]  # python-docx package

    doc = docx.Document(io.BytesIO(data))
    parts: list[str] = []

    # Paragraphs in document order.
    for p in doc.paragraphs:
        t = p.text.strip()
        if t:
            parts.append(t)

    # Tables — flatten cells. Doesn't try to preserve layout, just
    # gets the text out so DLP has something to inspect.
    table_count = 0
    for tbl in doc.tables:
        table_count += 1
        for row in tbl.rows:
            cells = [cell.text.strip() for cell in row.cells]
            cells = [c for c in cells if c]
            if cells:
                parts.append(" | ".join(cells))

    raw = "\n".join(parts)
    full_chars = len(raw)
    truncated = full_chars > MAX_EXTRACT_CHARS
    text = raw[:MAX_EXTRACT_CHARS] if truncated else raw
    if truncated:
        text += (
            f"\n\n[...truncated, full document was {full_chars:,} characters; "
            f"showing first {MAX_EXTRACT_CHARS:,}]"
        )

    bits = [f"DOCX, {len(doc.paragraphs)} paragraphs"]
    if table_count:
        bits.append(f"{table_count} tables")
    bits.append(f"{full_chars:,} chars")
    if truncated:
        bits.append("truncated")

    return ExtractResult(
        text=text,
        char_count=full_chars,
        truncated=truncated,
        source_kind="docx",
        summary=", ".join(bits),
    )


# ---------------------------------------------------------------------------
# XLSX — openpyxl
# ---------------------------------------------------------------------------

def _extract_xlsx(filename: str, data: bytes) -> ExtractResult:
    import openpyxl  # type: ignore[import-not-found]

    # data_only=True → formula cells return their cached value, not
    # the formula string. read_only=True → don't try to load styles
    # and other heavy chrome. We just want the text.
    wb = openpyxl.load_workbook(
        io.BytesIO(data), data_only=True, read_only=True,
    )

    parts: list[str] = []
    sheet_count = 0
    row_count = 0
    for sheet_name in wb.sheetnames:
        sheet_count += 1
        sh = wb[sheet_name]
        parts.append(f"--- Sheet: {sheet_name} ---")
        for row in sh.iter_rows(values_only=True):
            row_count += 1
            cells = [
                str(v) for v in row
                if v is not None and str(v).strip()
            ]
            if cells:
                parts.append("\t".join(cells))

    raw = "\n".join(parts)
    full_chars = len(raw)
    truncated = full_chars > MAX_EXTRACT_CHARS
    text = raw[:MAX_EXTRACT_CHARS] if truncated else raw
    if truncated:
        text += (
            f"\n\n[...truncated, full workbook was {full_chars:,} characters; "
            f"showing first {MAX_EXTRACT_CHARS:,}]"
        )

    bits = [
        f"XLSX, {sheet_count} sheet{'s' if sheet_count != 1 else ''}",
        f"~{row_count} rows",
        f"{full_chars:,} chars",
    ]
    if truncated:
        bits.append("truncated")

    return ExtractResult(
        text=text,
        char_count=full_chars,
        truncated=truncated,
        source_kind="xlsx",
        summary=", ".join(bits),
    )


# ---------------------------------------------------------------------------
# Public dispatcher
# ---------------------------------------------------------------------------

def extract(filename: str, content: bytes) -> ExtractResult:
    """Extract text from a single uploaded file.

    Raises ``ValueError`` for unsupported extensions. Library-level
    exceptions during extraction are NOT swallowed — the web layer
    catches them and returns a 422 with the error message, so the
    operator sees what went wrong (e.g. corrupt PDF, encrypted DOCX).
    """
    ext = Path(filename).suffix.lower()

    if ext in INLINE_TEXT_EXTS:
        return _extract_inline_text(filename, content)
    if ext == ".pdf":
        return _extract_pdf(filename, content)
    if ext == ".docx":
        return _extract_docx(filename, content)
    if ext == ".xlsx":
        return _extract_xlsx(filename, content)

    raise ValueError(
        f"unsupported file type: {ext!r}. Supported types: "
        f"{', '.join(sorted(SUPPORTED_EXTS))}"
    )
